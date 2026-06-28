# -*- coding: utf-8 -*-
import tkinter as tk
from tkinter import ttk, messagebox, scrolledtext
import threading
import time
import json
import urllib.request
import os
from docx import Document
from docx.shared import Pt, RGBColor

class MultiAgentGUIApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Pixel Agents - GPT Multi-Agent 한국어 협업 엔진")
        self.root.geometry("1020x720")
        self.root.configure(bg="#0B132B") # 프리미엄 다크 네이비 배경
        
        # 폰트 설정
        self.font_title = ("Malgun Gothic", 16, "bold")
        self.font_subtitle = ("Malgun Gothic", 11, "bold")
        self.font_body = ("Malgun Gothic", 10)
        self.font_log = ("Consolas", 10)
        
        # 테마 색상 설정
        self.color_bg = "#0B132B"
        self.color_card = "#1C2541"
        self.color_accent = "#48CAE4"
        self.color_accent_hover = "#00B4D8"
        self.color_text = "#FFFFFF"
        self.color_muted = "#8D99AE"
        self.color_status_active = "#2ECC71"
        self.color_status_idle = "#F1C40F"
        
        # 가상 에이전트 역할 데이터베이스
        self.agent_roles = [
            {"title": "👑 Project Manager", "role": "PM (조율 및 설계)", "color": "#FF6B6B", 
             "sys": "당신은 협업 에이전트 시스템의 PM(Project Manager)입니다. 사용자의 원래 요청을 꼼꼼히 분석하여 다음 단계 연구원(Researcher)에게 지시할 작업 분석 요구서와 수행 가이드를 한국어로 정중하게 작성하십시오."},
            {"title": "🔍 Info Analyst", "role": "Researcher (자료 분석)", "color": "#4D96FF",
             "sys": "당신은 연구원(Researcher) 에이전트입니다. 앞서 PM 에이전트가 지시한 분석 요구서를 해독하여, 해당 주제에 적합한 구체적인 데이터 조사 결과, 이론적 분석 내용, 핵심 요약 포인트를 한국어로 깊이 있게 상세 리포트 형식으로 작성하십시오."},
            {"title": "💻 Software Engineer", "role": "Developer (설계 및 구현)", "color": "#6BCB77",
             "sys": "당신은 개발자 및 작가(Developer/Writer) 에이전트입니다. 연구원(Researcher)의 조사 리포트를 바탕으로 원래 사용자가 물어본 질문에 대해 완벽하게 적용 가능한 상세 구조안이나 솔루션(구현 코드, 텍스트 본문 등)을 한국어로 친절하게 작성하십시오."},
            {"title": "🧪 Quality Reviewer", "role": "Reviewer (검증 및 보완)", "color": "#FFD93D",
             "sys": "당신은 품질 관리자(Reviewer) 에이전트입니다. 개발자가 구현해 낸 최종본 소스를 검사하여 가독성, 예외 처리, 타입 오류 등을 체크하고, 이를 더욱 세련되게 보완 및 개선한 한글 최종 결과물을 도출하십시오."},
            {"title": "📄 Technical Writer", "role": "Writer (최종 보고서 작성)", "color": "#B388FF",
             "sys": "당신은 테크니컬 라이터 에이전트입니다. 지금까지 모든 에이전트들의 협업 이력(PM 기획, 분석, 구현, 검수)을 종합하여 개요, 본문, 개선점을 일목요연하게 갖춘 한글 종합 최종 보고서 형태로 정돈하여 완성하십시오."}
        ]
        
        self.running_simulation = False
        self.setup_ui()
        
    def setup_ui(self):
        style = ttk.Style()
        style.theme_use('clam')
        style.configure('TFrame', background=self.color_bg)
        
        # 전체 그리드 구조 세팅
        self.root.columnconfigure(0, weight=4) # 컨트롤 패널
        self.root.columnconfigure(1, weight=6) # 로그 및 시각화 패널
        self.root.rowconfigure(0, weight=1)
        
        # ── LEFT: 컨트롤 패널 ──
        left_panel = tk.Frame(self.root, bg=self.color_card, padx=20, pady=20)
        left_panel.grid(row=0, column=0, sticky="nsew", padx=10, pady=10)
        
        # 제목
        lbl_title = tk.Label(left_panel, text="Pixel Agents", fg=self.color_accent, bg=self.color_card, font=("Malgun Gothic", 20, "bold"))
        lbl_title.pack(anchor="w", pady=(0, 5))
        
        lbl_desc = tk.Label(left_panel, text="GPT 연동 한국어 협업 엔진 v1.5", fg=self.color_muted, bg=self.color_card, font=self.font_body)
        lbl_desc.pack(anchor="w", pady=(0, 20))
        
        self.add_divider(left_panel)
        
        # API Key 입력 영역
        lbl_api = tk.Label(left_panel, text="OpenAI API Key (입력 시 실시간 GPT 연동)", fg=self.color_text, bg=self.color_card, font=self.font_subtitle)
        lbl_api.pack(anchor="w", pady=(10, 5))
        self.entry_api = tk.Entry(left_panel, show="*", bg="#0B132B", fg=self.color_text, insertbackground="white", borderwidth=1, relief="solid", font=self.font_body)
        self.entry_api.pack(fill="x", pady=(0, 5))
        
        lbl_api_tip = tk.Label(left_panel, text="※ 미입력 시에는 유려한 로컬 시뮬레이션 모드로 작동합니다.", fg=self.color_muted, bg=self.color_card, font=("Malgun Gothic", 9))
        lbl_api_tip.pack(anchor="w", pady=(0, 15))
        
        # 멀티 에이전트 개수 설정 영역
        lbl_num = tk.Label(left_panel, text="가동할 멀티 에이전트 개수 설정", fg=self.color_text, bg=self.color_card, font=self.font_subtitle)
        lbl_num.pack(anchor="w", pady=(5, 5))
        
        # 3~5개 추천 문구
        lbl_recommend = tk.Label(left_panel, text="💡 추천 개수: 3개 ~ 5개 (API 제한 및 컨텍스트 비용 최적)", fg=self.color_accent, bg=self.color_card, font=("Malgun Gothic", 9))
        lbl_recommend.pack(anchor="w", pady=(0, 5))
        
        self.slider_agents = tk.Scale(left_panel, from_=1, to=5, orient="horizontal", bg=self.color_card, fg=self.color_text, highlightbackground=self.color_card, activebackground=self.color_accent, troughcolor="#0B132B", font=self.font_body)
        self.slider_agents.set(3) # 기본값 3개
        self.slider_agents.pack(fill="x", pady=(0, 15))
        
        # 질문 입력 영역
        lbl_prompt = tk.Label(left_panel, text="한국어 지시사항 / 질문 입력", fg=self.color_text, bg=self.color_card, font=self.font_subtitle)
        lbl_prompt.pack(anchor="w", pady=(5, 5))
        
        self.text_prompt = scrolledtext.ScrolledText(left_panel, height=8, bg="#0B132B", fg=self.color_text, insertbackground="white", borderwidth=1, relief="solid", font=self.font_body)
        self.text_prompt.pack(fill="x", pady=(0, 20))
        self.text_prompt.insert(tk.END, "행안부 고급인증과정 6주차 과제에 필요한 다중 에이전트의 효율적인 협업 흐름을 설계해줘.")
        
        # 실행 단추
        self.btn_run = tk.Button(left_panel, text="🚀 멀티 에이전트 협업 실행하기", bg=self.color_accent, fg=self.color_bg, activebackground=self.color_accent_hover, activeforeground=self.color_bg, font=self.font_subtitle, relief="flat", command=self.start_simulation)
        self.btn_run.pack(fill="x", pady=(5, 10))
        
        # DOCX 저장 단추 (비활성화 상태로 시작)
        self.btn_save = tk.Button(left_panel, text="💾 최종 협업 보고서 DOCX 저장", bg="#2ECC71", fg=self.color_bg, activebackground="#27AE60", activeforeground=self.color_bg, font=self.font_subtitle, relief="flat", state="disabled", command=self.save_collaboration_report)
        self.btn_save.pack(fill="x", pady=(5, 0))
        
        # ── RIGHT: 로그 및 모니터링 패널 ──
        right_panel = tk.Frame(self.root, bg=self.color_bg, padx=10, pady=10)
        right_panel.grid(row=0, column=1, sticky="nsew")
        right_panel.columnconfigure(0, weight=1)
        right_panel.rowconfigure(0, weight=3) # 에이전트 캐릭터 스폰 오피스
        right_panel.rowconfigure(1, weight=7) # 대화 로그
        
        # 에이전트 오피스 시각화 프레임
        self.office_frame = tk.LabelFrame(right_panel, text="Pixel Agents - 가상 오피스 현황", fg=self.color_accent, bg=self.color_card, font=self.font_subtitle, padx=10, pady=10, borderwidth=1, relief="solid")
        self.office_frame.grid(row=0, column=0, sticky="nsew", pady=(0, 10))
        
        self.agent_slots = []
        
        # 대화 및 협업 로그 프레임
        log_frame = tk.LabelFrame(right_panel, text="에이전트 한국어 협업 로그", fg=self.color_accent, bg=self.color_card, font=self.font_subtitle, padx=10, pady=10, borderwidth=1, relief="solid")
        log_frame.grid(row=1, column=0, sticky="nsew")
        
        self.log_area = scrolledtext.ScrolledText(log_frame, bg="#0B132B", fg="#E0E0E0", insertbackground="white", borderwidth=0, font=self.font_log)
        self.log_area.pack(fill="both", expand=True)
        self.log_area.insert(tk.END, "[대기 중] 좌측 패널에서 에이전트 수와 질문을 입력하고 시작 단추를 클릭하세요.\n")
        
    def add_divider(self, parent):
        divider = tk.Frame(parent, height=1, bg=self.color_muted, bd=0)
        divider.pack(fill="x", pady=10)
        
    def start_simulation(self):
        if self.running_simulation:
            messagebox.showwarning("경고", "이미 멀티 에이전트 프로세스가 구동 중입니다.")
            return
        
        prompt = self.text_prompt.get("1.0", tk.END).strip()
        if not prompt:
            messagebox.showwarning("경고", "질문이나 지시사항을 입력해 주세요.")
            return
            
        self.running_simulation = True
        self.btn_run.config(state="disabled", bg=self.color_muted)
        self.btn_save.config(state="disabled")
        
        # 백그라운드 스레드에서 실제 / 가상 체인 생성 시작
        threading.Thread(target=self.run_agents_workflow, args=(prompt,), daemon=True).start()
        
    def call_gpt_api(self, api_key, system_prompt, prompt_text):
        url = "https://api.openai.com/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        data = {
            "model": "gpt-4o",
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": prompt_text}
            ],
            "temperature": 0.7
        }
        
        req = urllib.request.Request(url, data=json.dumps(data).encode("utf-8"), headers=headers, method="POST")
        try:
            with urllib.request.urlopen(req, timeout=45) as response:
                res_data = json.loads(response.read().decode("utf-8"))
                return res_data["choices"][0]["message"]["content"].strip()
        except Exception as e:
            raise Exception(f"GPT API 호출 실패: {str(e)}")

    def run_agents_workflow(self, prompt):
        agent_count = self.slider_agents.get()
        api_key = self.entry_api.get().strip()
        is_live_api = len(api_key) > 0
        
        # 1. 오피스에 에이전트 스폰 및 렌더링
        self.root.after(0, self.spawn_office_agents, agent_count)
        
        self.log_clear()
        self.log_write("🔔 멀티 에이전트 시스템이 초기화되었습니다.\n")
        self.log_write(f"👥 설정된 에이전트 수: {agent_count}개\n")
        self.log_write(f"🌐 연동 모드: {'[실시간 OpenAI GPT 연동 모드]' if is_live_api else '[로컬 고성능 시뮬레이션 모드]'}\n")
        self.log_write(f"💬 사용자 요청: \"{prompt}\"\n")
        self.log_write("─" * 60 + "\n")
        
        time.sleep(1.0)
        
        selected_roles = self.agent_roles[:agent_count]
        current_input = prompt
        
        for i, agent in enumerate(selected_roles):
            # 에이전트 상태를 Active(활동중)로 변경
            self.root.after(0, self.update_agent_status, i, "Active 💻", self.color_status_active)
            self.log_write(f"\n[{agent['title']}] ({agent['role']}) 활성화 중...\n")
            
            output_text = ""
            
            if is_live_api:
                # ── 실제 OpenAI GPT-4o API 호출 및 릴레이 체이닝 ──
                try:
                    self.log_write(f"⚡ [API]: GPT-4o 연동 호출 중...\n")
                    output_text = self.call_gpt_api(api_key, agent["sys"], current_input)
                except Exception as e:
                    self.log_write(f"❌ [에러 발생]: {str(e)}\n시뮬레이션 모드로 전환하여 작업을 계속 진행합니다.\n")
                    is_live_api = False # 실패 시 시뮬레이션 모드로 전환
                    
            if not is_live_api:
                # ── 로컬 Mock 시뮬레이션 데이터 생성 ──
                time.sleep(2.0)
                if agent["role"].startswith("PM"):
                    output_text = (
                        f"[PM 기획서]\n"
                        f"1. 목표: 행안부 고급인증과정 6주차 다중 에이전트 협업 아키텍처 설계\n"
                        f"2. 핵심 요구사항: 에이전트 간 결합도 완화 및 비동기 JSONL 스캔 기능 최적화\n"
                        f"3. Researcher 지시사항: CLI 훅 포트 통신 구조 및 비동기 감지 실패 케이스를 조사하시오."
                    )
                elif "Researcher" in agent["role"]:
                    output_text = (
                        f"[Researcher 리포트]\n"
                        f"- 조사 결과: GPT 기반 API 통신에서는 레이턴시 제어가 핵심이므로 로컬 파이프라인 캐싱 도입 권장\n"
                        f"- 분석 테마: 비동기 병렬 스폰 시 API 제한 극복을 위해 릴레이 큐(Queue) 패턴 적용 필요 확인\n"
                        f"- 해결 제안: Developer 에이전트는 Queue 처리를 지원하는 TS 기반 비동기 핸들러 코드를 구현하라."
                    )
                elif "Developer" in agent["role"]:
                    output_text = (
                        f"[Developer 설계안]\n"
                        f"- 아키텍처: Queue 기반 Multi-Agent 스케줄러 구현\n"
                        f"- 구현 코드(TS):\n"
                        f"  class AgentQueue {{ private queue: Task[] = []; async push(t: Task) {{ ... }} }}\n"
                        f"- 특이사항: 이 구조를 통해 3~5개 다중 에이전트가 호출 제한 없이 부드럽게 병렬 동작하도록 개선 완료."
                    )
                elif "Reviewer" in agent["role"]:
                    output_text = (
                        f"[Reviewer 검증 본문]\n"
                        f"- 타입 안정성 검수: TypeScript tsconfig 기준 무결성 통과\n"
                        f"- 예외 처리 검토: API 호출 타임아웃 발생 시 fallback 핸들러 바인딩 코드 추가 보완\n"
                        f"- 보완 완료: 예외 예방율 98.7% 향상."
                    )
                else:
                    output_text = (
                        f"[Writer 종합 보고서]\n"
                        f"본 행안부 6주차 과제를 위한 다중 에이전트 시스템은 PM의 스케줄링, Researcher의 분석 제안,\n"
                        f"Developer의 비동기 Queue 설계 및 Reviewer의 안전성 통과를 거쳐 고신뢰도 멀티 에이전트 솔루션으로\n"
                        f"설계 및 최종 완성되었습니다."
                    )
            
            # 결과 기록 및 출력
            self.log_write(f"📝 [결과물]:\n{output_text}\n")
            current_input = f"이전 에이전트의 결과물은 다음과 같습니다:\n{output_text}\n이를 참조하여 당신의 역할을 수행하십시오."
            
            # 에이전트 상태를 Waiting(대기)로 변경
            self.root.after(0, self.update_agent_status, i, "Waiting 💤", self.color_status_idle)
            time.sleep(1.2)
            
        self.log_write("\n" + "─" * 60 + "\n")
        self.log_write("🎉 모든 에이전트의 실시간 협업 체인이 안전하게 마쳤습니다!\n")
        
        self.running_simulation = False
        self.root.after(0, self.enable_save_button)
        
    def spawn_office_agents(self, count):
        for widget in self.office_frame.winfo_children():
            widget.destroy()
            
        self.agent_slots = []
        
        for i in range(count):
            agent = self.agent_roles[i]
            
            card = tk.Frame(self.office_frame, bg=self.color_bg, borderwidth=1, relief="solid", padx=8, pady=8)
            card.pack(side="left", padx=5, fill="both", expand=True)
            
            lbl_icon = tk.Label(card, text="🤖", font=("Segoe UI Emoji", 20), fg=self.color_accent, bg=self.color_bg)
            lbl_icon.pack(anchor="center")
            
            lbl_name = tk.Label(card, text=agent["title"], fg=agent["color"], bg=self.color_bg, font=self.font_body, wraplength=100, justify="center")
            lbl_name.pack(anchor="center", pady=(5, 0))
            
            lbl_role = tk.Label(card, text=agent["role"], fg=self.color_muted, bg=self.color_bg, font=("Malgun Gothic", 9))
            lbl_role.pack(anchor="center")
            
            lbl_status = tk.Label(card, text="Waiting 💤", fg=self.color_status_idle, bg=self.color_bg, font=("Malgun Gothic", 9, "bold"))
            lbl_status.pack(anchor="center", pady=(5, 0))
            
            self.agent_slots.append(lbl_status)
            
    def update_agent_status(self, index, text, color):
        if index < len(self.agent_slots):
            self.agent_slots[index].config(text=text, fg=color)
            
    def enable_save_button(self):
        self.btn_run.config(state="normal", bg=self.color_accent)
        self.btn_save.config(state="normal")
        
    def log_clear(self):
        self.log_area.delete("1.0", tk.END)
        
    def log_write(self, text):
        self.log_area.insert(tk.END, text)
        self.log_area.see(tk.END)
        
    def save_collaboration_report(self):
        try:
            doc = Document()
            style = doc.styles['Normal']
            style.font.name = 'Malgun Gothic'
            style.font.size = Pt(11)
            
            title = doc.add_paragraph()
            title.alignment = 1 # Center
            run_t = title.add_run("Pixel Agents - 멀티 에이전트 한국어 협업 리포트 (GPT)")
            run_t.font.size = Pt(18)
            run_t.font.bold = True
            run_t.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
            
            doc.add_paragraph(f"생성일자: {time.strftime('%Y-%m-%d %H:%M:%S')}\n질문내용: {self.text_prompt.get('1.0', tk.END).strip()}")
            doc.add_paragraph("─" * 50)
            
            doc.add_heading("1. 협업 세부 내용", level=2)
            doc.add_paragraph(self.log_area.get("1.0", tk.END))
            
            output_file = "Pixel_Agents_Collaboration.docx"
            doc.save(output_file)
            
            abs_path = os.path.abspath(output_file)
            messagebox.showinfo("성공", f"협업 리포트가 성공적으로 저장되었습니다!\n경로: {abs_path}")
        except Exception as e:
            messagebox.showerror("에러", f"파일 저장 중 에러가 발생했습니다: {str(e)}")

def main():
    root = tk.Tk()
    app = MultiAgentGUIApp(root)
    root.mainloop()

if __name__ == '__main__':
    main()
