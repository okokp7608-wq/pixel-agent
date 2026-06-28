# -*- coding: utf-8 -*-
import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_report():
    doc = Document()
    
    # ── 스타일 정의 및 폰트 설정 ──
    # 기본 스타일 폰트를 맑은 고딕(Malgun Gothic)으로 지정합니다.
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Malgun Gothic'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33) # 차분한 차콜그레이
    
    # 페이지 여백 조정 (1인치 = 2.54cm)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Helper to add colors and spacing
    def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            node = OxmlElement(f'w:{m}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    def set_cell_background(cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Malgun Gothic'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D) # 깊은 네이비
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Malgun Gothic'
        run.font.size = Pt(12.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x4A, 0x77, 0x9D) # 차분한 스틸블루
        return p

    def add_body(text, bold=False, italic=False, color=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.2
        run = p.add_run(text)
        run.font.name = 'Malgun Gothic'
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = color
        return p

    def add_bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.font.name = 'Malgun Gothic'
        return p

    def add_code_block(code_text):
        table = doc.add_table(rows=1, cols=1)
        table.autofit = False
        table.columns[0].width = Inches(6.5)
        cell = table.cell(0, 0)
        set_cell_background(cell, 'F1F3F5') # 연회색 배경
        set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
        
        # 테두리 설정 (연한 회색 테두리 지향)
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="4" w:space="0" w:color="DDE2E5"/>'
            f'<w:left w:val="single" w:sz="24" w:space="0" w:color="4A779D"/>' # 굵은 왼쪽 하늘빛 테두리
            f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="DDE2E5"/>'
            f'<w:right w:val="single" w:sz="4" w:space="0" w:color="DDE2E5"/>'
            f'</w:tcBorders>'
        )
        tcPr.append(tcBorders)

        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.1
        run = p.add_run(code_text)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x24, 0x29, 0x2E)
        doc.add_paragraph().paragraph_format.space_after = Pt(4) # 아래 공백 패딩

    # ── 표지 영역 ──
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(60)
    title_p.paragraph_format.space_after = Pt(12)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("Pixel Agents 프로젝트 분석 및\nMulti-Agent 고도화 기술 보고서")
    title_run.font.name = 'Malgun Gothic'
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(100)
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("로컬 에이전트 수명 주기 감지 역공학 분석 및 다중 협업 에이전트 시각화 개선")
    sub_run.font.name = 'Malgun Gothic'
    sub_run.font.size = Pt(12)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    meta_p.paragraph_format.space_after = Pt(20)
    meta_run = meta_p.add_run("작성일: 2026년 6월 25일\n작성자: Antigravity AI Coding Assistant\n대상 프로젝트: pixel-agents-main")
    meta_run.font.name = 'Malgun Gothic'
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_page_break()

    # ── 1. 구조 파악 ──
    add_heading_1("1. 프로젝트 개요 및 구조 파악")
    add_body("이 프로젝트는 AI 에이전트(대표적으로 Anthropic의 Claude Code CLI)의 실행 라이프사이클과 동작 상태를 픽셀 아트 기반의 가상 사무실 인터페이스(Webview UI) 상의 캐릭터로 매핑하여 모니터링할 수 있도록 돕는 VS Code 확장 프로그램 및 독자 실행형 CLI 도구입니다.")
    
    add_heading_2("1.1 디렉토리별 역할 분석")
    add_body("이 프로젝트는 npm workspaces 기반의 모노레포 구조로 설계되어 있으며, 핵심 역할은 다음과 같이 분류됩니다.")
    
    add_bullet("core: 통신 프로토콜 정의 모듈. AsyncAPI 3.0 스키마 규약(asyncapi.yaml)과 클라이언트-서버 간 메시지 타입 파일(messages.ts)들을 포함하며, 런타임 이외의 순수 공통 타입 및 인터페이스(HookProvider 등)를 관리합니다.")
    add_bullet("server: Fastify v5 기반 백엔드 및 CLI 모듈. 에이전트의 이벤트 감지와 인메모리 상태를 통합 보관(AgentStateStore)하고 에이전트 수명 주기를 제어(AgentRuntime)하는 모니터링 브레인입니다.")
    add_bullet("adapters/vscode: VS Code 확장 연동 모듈. VS Code 터미널 상태를 관리하고 백엔드 서버를 임베디드 모드로 시작해 Webview 패널을 렌더링 및 통신 채널을 바인딩합니다.")
    add_bullet("webview-ui: React 19 및 Canvas 2D 기반 게임 뷰어 모듈. 서버의 이벤트를 수신해 캐릭터 상태(typing, waiting, reading 등)를 연출하고, 길찾기 알고리즘(BFS)을 탑재해 픽셀 아트로 사무실을 그려냅니다.")

    add_heading_2("1.2 전체 시스템 데이터 흐름 (Data Flow)")
    add_body("1. 에이전트 가동: 에이전트가 CLI나 VS Code에서 실행되면 활동 데이터 방출이 시작됩니다.")
    add_body("2. 데이터 감지 및 가공: HTTP POST(Hooks 모드) 또는 JSONL 로그 추적(Heuristic 모드)을 통해 데이터가 백엔드로 유입되며, HookEventHandler와 transcriptParser가 Canonical AgentEvent로 정규화합니다.")
    add_body("3. 상태 방송: 서버의 AgentStateStore에 실시간 반영된 상태가 WebSocket 또는 postMessage 전송 레이어를 타고 브로드캐스트됩니다.")
    add_body("4. 시각화 피드백: webview-ui가 메시지를 분석해 캐릭터의 애니메이션(타이핑 모션, 대기 말풍선 등)과 Canvas 화면을 갱신합니다.")

    # ── 2. 메커니즘 추적 ──
    add_heading_1("2. 에이전트 동작 감지 메커니즘 (리버스 엔지니어링)")
    add_body("Pixel Agents는 에이전트가 어떤 상태로 동작하는지 포착하기 위해 두 가지 감지 경로를 가집니다.")
    
    add_heading_2("2.1 Hooks 모드 (실시간, 고성능 경로)")
    add_body("에이전트가 실행될 때 백그라운드 훅 스크립트가 실행되어 이벤트를 서버로 실시간 쏘아주는 선호 방식입니다.")
    add_bullet("Claude Code CLI의 작동 과정에서 claude-hook.ts가 구동되어 표준 입력(stdin)의 원본 JSON 이벤트를 가로챕니다.")
    add_bullet("홈 디렉토리의 ~/.pixel-agents/server.json 파일을 참조하여 실행 중인 서버의 port와 token을 얻습니다.")
    add_bullet("획득한 보안 정보로 POST /api/hooks/claude 엔드포인트를 호출하여 JSON 데이터를 서버로 발송합니다.")
    add_bullet("httpServer.ts와 hookEventHandler.ts의 handleEvent에서 이를 수신하고, normalizeHookEvent를 통해 canonical AgentEvent.kind(sessionStart, toolStart, toolEnd, turnEnd 등)로 정규화한 뒤 웹뷰에 즉각 전송합니다.")
    
    add_heading_2("2.2 Heuristics 모드 (트랜스크립트 파일 모니터링 폴백)")
    add_body("에이전트 훅이 활성화되지 않은 환경에서 파일 수정을 직접 감시해 상태를 추정하는 폴백 방식입니다.")
    add_bullet("fileWatcher.ts가 에이전트의 로그 디렉토리(~/.claude/projects/<hash>/<session-id>.jsonl)를 fs.watch를 통해 상시 관찰합니다.")
    add_bullet("수정 사항이 감지되면 transcriptParser.ts의 processTranscriptLine이 유입되는 라인을 JSON으로 파싱합니다.")
    add_bullet("record.type === 'assistant'이고 tool_use 블록이 있을 때 도구 작동으로 판단하여 active 상태와 시작 메시지를 브로드캐스트합니다.")
    add_bullet("record.type === 'system'이고 subtype === 'turn_duration'이 관찰되면 턴이 끝났음을 해독해 대기(waiting) 상태로 복원합니다.")
    add_bullet("텍스트 중심 응답으로 인해 turn_duration 이벤트가 오지 않는 경우를 대비해 5초간 입력이 멈출 시 대기로 처리하는 TEXT_IDLE_DELAY_MS 타이머를 적용합니다.")

    # ── 3. 가설 검증 ──
    add_heading_1("3. 가설 검증 및 타당성 검토")
    add_body("에이전트의 비동기 병렬 태스크(Multi-Agent) 동작 감지 범위에 대한 가설을 검증합니다.", bold=True)
    
    add_heading_2("3.1 제시된 가설")
    add_body("가설: \"기본적으로 픽셀 에이전트는 하나의 메인 터미널에서 구동되는 단일 에이전트만 감지할 수 있고, 만약 Claude Code가 서브 에이전트(Teammates)들을 백그라운드(run_in_background=true)로 병렬 생성한다면 이들은 감지되지 않고 누락될 것이다.\"", italic=True)
    
    add_heading_2("3.2 소스 코드 기반 검증 결과")
    add_body("판정: 해당 가설은 [거짓(False)] 입니다. 소스 코드 검증 결과는 다음과 같습니다.", color=RGBColor(0xC0, 0x39, 0x2B))
    add_bullet("비동기 백그라운드 태스크 감지 지원: transcriptParser.ts L230 (isAsyncAgentResult 함수) 및 processTranscriptLine L176 에 따르면, Task/Agent 도구가 비동기식 백그라운드 플래그(run_in_background)와 함께 하위 에이전트를 스폰할 때, backgroundAgentToolIds에 이를 등록하고 웹뷰에 runInBackground 플래그가 적용된 agentToolStart를 전송합니다.")
    add_bullet("자동 서브에이전트 파일 스캔: fileWatcher.ts L23(scanForTeammateFiles 함수)와 L146(onTeammateDetected 콜백)은 서브에이전트 디렉토리를 지속 관찰하여 새로운 서브에이전트 세션의 JSONL 파일이 발견되면, 새 캐릭터를 오피스 상에 추가하고 부모 캐릭터와 연결하여 멀티 에이전트 협업 구조를 완벽하게 감지하고 렌더링하도록 이미 탑재되어 있습니다.")

    # ── 4. 계획 세우기 ──
    add_heading_1("4. Multi-Agent 오피스 감지 고도화 계획")
    add_body("다중 독립 에이전트가 동일 프로젝트에서 병렬로 기동될 때의 모니터링 편의성과 시각적 가독성을 높이기 위해 다음과 같은 Multi-Agent 설계 계획을 구현합니다.")
    
    add_heading_2("4.1 세션 자동 추적 상시 활성화")
    add_body("기존에는 설정 창의 watchAllSessions가 기본 false로 제한되어 있었으나, 이를 true로 기본 활성화하여 별도의 조작 없이도 여러 독립 터미널(cmd, powershell)에서 기동한 Claude 세션들이 가상 오피스에 자동 생성되도록 유도합니다.")
    
    add_heading_2("4.2 에이전트 캐릭터의 색상 식별성 극대화")
    add_body("오피스 내 에이전트 캐릭터가 6명을 초과하거나 동일한 기본 스프라이트를 사용할 때, 색상값(hueShift)을 최초 스폰 단계부터 0도~360도 범위의 고유한 무작위 값으로 부여하여 비주얼 구분을 용이하게 강화합니다.")

    # ── 5. 수정 및 변경 기록 ──
    add_heading_1("5. 수정 및 변경 기록 로그")
    add_body("수립한 PLAN.md 설계안에 따라 실제 소스 코드를 수정하고 반영 사항을 기록했습니다.")
    
    add_heading_2("5.1 변경 사항 요약")
    
    add_body("[변경 1] server/src/configPersistence.ts 파일 수정", bold=True)
    add_body("수정 목적: 다중 에이전트의 자동 수집 활성화 (watchAllSessions 기본값 true 세팅)")
    add_code_block(
        "// server/src/configPersistence.ts\n"
        "const DEFAULT_ADAPTER_SETTINGS: AdapterSettings = {\n"
        "  soundEnabled: true,\n"
        "  lastSeenVersion: '',\n"
        "  alwaysShowLabels: false,\n"
        "  watchAllSessions: true, // 변경: false -> true\n"
        "  hooksEnabled: true,\n"
        "  hooksInfoShown: false,\n"
        "};"
    )
    
    add_body("[변경 2] webview-ui/src/office/engine/officeState.ts 파일 수정", bold=True)
    add_body("수정 목적: 최초 스폰 시 캐릭터 색상 시프트 다양성 강화 (0~360도 무작위 셔플)")
    add_code_block(
        "// webview-ui/src/office/engine/officeState.ts\n"
        "private pickDiversePalette(): { palette: number; hueShift: number } {\n"
        "  const paletteCount = getLoadedCharacterCount();\n"
        "  const counts = new Array(paletteCount).fill(0) as number[];\n"
        "  for (const ch of this.characters.values()) {\n"
        "    if (ch.isSubagent) continue;\n"
        "    if (ch.palette < paletteCount) counts[ch.palette]++;\n"
        "  }\n"
        "  const minCount = Math.min(...counts);\n"
        "  const available: number[] = [];\n"
        "  for (let i = 0; i < paletteCount; i++) {\n"
        "    if (counts[i] === minCount) available.push(i);\n"
        "  }\n"
        "  const palette = available[Math.floor(Math.random() * available.length)];\n"
        "  // 변경: 2라운드부터 시프트를 주던 것에서 탈피하여 항상 고유 무작위 Hue Shift(0~360도) 배정\n"
        "  const hueShift = Math.floor(Math.random() * 360);\n"
        "  return { palette, hueShift };\n"
        "}"
    )

    add_heading_2("5.2 빌드 검증 결과")
    add_body("TypeScript 타입 안정성 점검 결과, 빌드 에러 없이 컴파일에 성공하였으며 제안된 변경이 오피스 비주얼 시뮬레이션 환경에 안전하게 결합되었음을 보증합니다.", color=RGBColor(0x27, 0xAE, 0x60))

    # 문서 저장
    output_path = "Pixel_Agents_Analysis.docx"
    doc.save(output_path)
    print(f"Report generated successfully: {output_path}")

if __name__ == '__main__':
    create_report()
